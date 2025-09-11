# ---------------- Imports ----------------
import os, re, shutil
from io import BytesIO

import streamlit as st
import google.generativeai as genai
import data_loader, vector_store
from config import GEMINI_API_KEY, CHAT_MODEL

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Optional: DOCX export
from docx import Document

# ---------------- Fonts & Model ----------------
pdfmetrics.registerFont(TTFont("DejaVuSans", "DejaVuSans.ttf"))
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel(CHAT_MODEL)

# ---------------- Config page & temp ----------------
st.set_page_config(page_title="Chat với tài liệu", layout="wide")
TEMP_DIR = "temp_uploads"
os.makedirs(TEMP_DIR, exist_ok=True)

# ---------------- Session state ----------------
st.session_state.setdefault("history", [])
st.session_state.setdefault("files", {})
st.session_state.setdefault("current_file", None)
st.session_state.setdefault("current_files", [])

# ---------------- Helpers ----------------
def unselect_file(): st.session_state["current_file"] = None; st.rerun()

def delete_file(fname):
    fdata = st.session_state.files.get(fname)
    if fdata:
        if os.path.exists(fdata["path"]): os.remove(fdata["path"])
        safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', os.path.splitext(fname)[0])
        faiss_folder = os.path.join(vector_store.FAISS_DIR, safe_name)
        if os.path.exists(faiss_folder): shutil.rmtree(faiss_folder)
        del st.session_state.files[fname]
        if st.session_state.get("current_file") == fname: st.session_state["current_file"] = None
    st.rerun()

def _clean_text_for_plain(s: str) -> str:
    if not s: return ""
    s = s.strip()
    s = re.sub(r'^[\u2022\*\-\s]+', '', s)
    s = re.sub(r'\*\*(.+?)\*\*', r'\1', s)
    s = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', s)
    s = re.sub(r':\*+', ':', s)
    return s

def _clean_text_for_pdf_html(s: str) -> str:
    if not s: return ""
    s = s.strip()
    s = re.sub(r'^[\u2022\*\-\s]+', '', s)
    s = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', s)
    s = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'<i>\1</i>', s)
    s = re.sub(r':\*+', ':', s)
    return s

# ---------------- Sidebar ----------------
with st.sidebar:
    st.markdown("## 📂 Quản lý tài liệu")
    uploaded_files = st.file_uploader(
        "Tải file (PDF, Word, Excel, Ảnh)...",
        type=["pdf","docx","xlsx","png","jpg","jpeg"],
        accept_multiple_files=True
    )

    if uploaded_files:
        for file in uploaded_files:
            if file.name in st.session_state.files: continue
            temp_path = os.path.join(TEMP_DIR, file.name)
            with open(temp_path, "wb") as f: f.write(file.read())
            text = data_loader.load_file(temp_path)
            chunks = vector_store.get_text_chunks(text) or [text]
            safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', os.path.splitext(file.name)[0])
            db_path = os.path.join(vector_store.FAISS_DIR, safe_name)
            os.makedirs(db_path, exist_ok=True)
            vector_store.save_vector_store(chunks, save_path=db_path)
            db = vector_store.load_vector_store(db_path)
            st.session_state.files[file.name] = {"file_obj": file, "path": temp_path, "text": text, "db": db, "db_path": db_path}

    if st.session_state.files:
        with st.expander("📂 Lịch sử file đã upload", expanded=True):
            file_names = list(st.session_state.files.keys())
            selected_files = st.multiselect(
                "Chọn file sử dụng:", file_names, default=st.session_state.get("current_files", [])
            )
            st.session_state["current_files"] = selected_files
            for fname in file_names:
                col1, col2 = st.columns([4,1])
                with col1:
                    st.markdown(f"**{fname}** 🟢" if fname in st.session_state["current_files"] else fname)
                with col2:
                    st.button("❌", key=f"del_{fname}", on_click=delete_file, args=(fname,))

# ---------------- Chat area ----------------
st.markdown("""
<style>
.stChatMessage.user {background:#d4f1ff; border-radius:12px; padding:10px; box-shadow:1px 1px 3px #ccc;}
.stChatMessage.assistant {background:#f0f0f0; border-radius:12px; padding:10px; box-shadow:1px 1px 3px #ccc;}
.typing {display:inline-block;width:1em;height:1em;border-radius:50%;background:#999;animation: blink 1.4s infinite; margin:0 2px;}
@keyframes blink {0%{opacity:.2;}20%{opacity:1;}100%{opacity:.2;}}
</style>
""", unsafe_allow_html=True)

st.title("💬 Chat với tài liệu")
for msg in st.session_state.history:
    role = msg["role"]
    with st.chat_message(role, avatar="🧑" if role=="user" else "🤖"):
        st.markdown(msg["content"])

if query := st.chat_input("Nhập câu hỏi..."):
    st.session_state.history.append({"role":"user","content":query})
    with st.chat_message("user", avatar="🧑"): st.markdown(query)
    with st.chat_message("assistant", avatar="🤖") as msg:
        typing_placeholder = st.empty()
        typing_placeholder.markdown('<span class="typing"></span>'*3, unsafe_allow_html=True)

        # Context & files
        file_names = list(st.session_state.files.keys())
        mentioned_files = [n for n in file_names if n.lower() in query.lower() or os.path.splitext(n.lower())[0] in query.lower()]
        if mentioned_files:
            target_files = [st.session_state.files[n] for n in mentioned_files]
            reading_info = ", ".join(mentioned_files)
        else:
            current_files = st.session_state.get("current_files", [])
            target_files = [st.session_state.files[n] for n in current_files] if current_files else []
            reading_info = ", ".join(current_files) if current_files else "không có file nào được chọn"

        typing_placeholder.markdown(f"📖 Bot đang đọc: **{reading_info}**<br><br>"+'<span class="typing"></span>'*3, unsafe_allow_html=True)

        context_texts = []
        for f in target_files:
            vs = vector_store.load_vector_store(f["db_path"])
            docs = vs.similarity_search(query, k=3)
            context_texts.extend([d.page_content for d in docs] if docs else [f["text"]])

        prompt = f"Ngữ cảnh:\n{'\n'.join(context_texts)}\n\nCâu hỏi: {query}\nTrả lời dựa trên ngữ cảnh trên."
        response = model.generate_content(prompt)

        typing_placeholder.empty()
        st.markdown(response.text)
        st.session_state.history.append({"role":"assistant","content":response.text})
        st.session_state.last_answer = response.text

# ---------------- Export ----------------
if "last_answer" in st.session_state and st.session_state.last_answer:
    with st.expander("📥 Xuất câu trả lời", expanded=True):
        export_format = st.radio("Chọn định dạng:", ["PDF","DOCX","TXT","Excel (code)"], horizontal=True)
        if export_format == "PDF":
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            style = ParagraphStyle(name="Vietnamese", fontName="DejaVuSans", fontSize=11, leading=15, spaceAfter=8)
            story, bullets = [], []
            for raw in st.session_state.last_answer.splitlines():
                line = raw.strip()
                if not line:
                    if bullets:
                        story.append(ListFlowable([ListItem(Paragraph(b,style)) for b in bullets], bulletType="bullet"))
                        bullets=[]
                    story.append(Spacer(1,8)); continue
                if re.match(r'^[\u2022\*\-\s]+', raw) or re.match(r'^\d+\.\s', line):
                    bullets.append(_clean_text_for_pdf_html(raw)); continue
                if bullets:
                    story.append(ListFlowable([ListItem(Paragraph(b,style)) for b in bullets], bulletType="bullet"))
                    bullets=[]
                story.append(Paragraph(_clean_text_for_pdf_html(line), style))
            if bullets: story.append(ListFlowable([ListItem(Paragraph(b,style)) for b in bullets], bulletType="bullet"))
            doc.build(story); buffer.seek(0)
            st.download_button("⬇️ Tải về PDF", data=buffer, file_name="chatbot_output.pdf", mime="application/pdf")

        elif export_format == "DOCX":
            buffer = BytesIO()
            doc = Document()
            def _add_docx_paragraph_with_md_runs(paragraph_obj, text):
                for tok in re.split(r'(\*\*.*?\*\*|\*.*?\*)', text):
                    if not tok: continue
                    if tok.startswith("**") and tok.endswith("**") and len(tok)>=4: paragraph_obj.add_run(tok[2:-2]).bold=True
                    elif tok.startswith("*") and tok.endswith("*") and len(tok)>=2: paragraph_obj.add_run(tok[1:-1]).italic=True
                    else: paragraph_obj.add_run(tok)
            for raw in st.session_state.last_answer.splitlines():
                line = raw.strip()
                if not line: doc.add_paragraph(""); continue
                if re.match(r'^[\u2022\*\-\s]+', raw):
                    p = doc.add_paragraph(style="List Bullet"); _add_docx_paragraph_with_md_runs(p,_clean_text_for_plain(raw)); continue
                if re.match(r'^\d+\.\s', line):
                    p = doc.add_paragraph(style="List Number"); _add_docx_paragraph_with_md_runs(p,_clean_text_for_plain(line)); continue
                p = doc.add_paragraph(); _add_docx_paragraph_with_md_runs(p,_clean_text_for_plain(line))
            doc.save(buffer); buffer.seek(0)
            st.download_button("⬇️ Tải về DOCX", data=buffer, file_name="chatbot_output.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        elif export_format == "TXT":
            cleaned_lines = [_clean_text_for_plain(l) if l.strip() else "" for l in st.session_state.last_answer.splitlines()]
            st.download_button("⬇️ Tải về TXT", data="\n".join(cleaned_lines), file_name="chatbot_output.txt", mime="text/plain")

        elif export_format == "Excel (code)":
            st.markdown("### 💡 Code xuất Excel (chạy local)")
            example_code = f'''
import pandas as pd
lines = [line.strip() for line in """{st.session_state.last_answer}""".split("\\n") if line.strip()]
df = pd.DataFrame({{"Nội dung": lines}})
df.to_excel("chatbot_output.xlsx", index=False, sheet_name="Chatbot Output")
print("✅ Đã lưu chatbot_output.xlsx")
            '''
            st.code(example_code, language="python")